import os
import docx
import torch
import shutil
import hashlib
import numpy as np
import unstructured
from typing import List
from pathlib import Path
import concurrent.futures
from dotenv import load_dotenv
from AI_ChatBot.logging import logger
from AI_ChatBot.utils.common import createDir
from AI_ChatBot.entity import DataTransformationParams,DataTransformationConfig

from langchain.schema import Document
from langchain_huggingface.embeddings import HuggingFaceEmbeddings
from langchain_openai.embeddings import OpenAIEmbeddings
from langchain_experimental.text_splitter import SemanticChunker
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_huggingface.embeddings.huggingface_endpoint import HuggingFaceEndpointEmbeddings
from langchain_community.document_loaders import UnstructuredPDFLoader,TextLoader,UnstructuredWordDocumentLoader


class DataTransformation:
    def __init__(self,config:DataTransformationConfig,params:DataTransformationParams,max_workers: int = None):
      self.config = config
      self.params = params
      load_dotenv()
      self.api_key = os.getenv("HUGGINGFACEHUB_API_TOKEN")
      self.api = os.getenv("OPENAI_API_KEY")
      if self.api_key:
        logger.info("Using Huggingface embeddings")
        self.embeddings = HuggingFaceEndpointEmbeddings(
          model="sentence-transformers/all-MiniLM-L6-v2",
          task="feature-extraction",          # required for embedding models
          huggingfacehub_api_token=self.api_key
      )
      elif self.api:
        logger.info("Using OpenAI embeddings")
        self.embeddings = OpenAIEmbeddings(api_key=self.api,model="text-embedding-3-small")
      else:
         logger.error("No Embedding Model")

      self.max_workers = max_workers or os.cpu_count() or 4
      self.recursive_splitter = RecursiveCharacterTextSplitter(
          chunk_size=self.params.recursive.chunk_size,
          chunk_overlap=self.params.recursive.chunk_overlap
      )
      self.chunker = SemanticChunker(
              embeddings=self.embeddings,
              breakpoint_threshold_type=self.params.semantic.breakpoint_threshold_type,
              breakpoint_threshold_amount=self.params.semantic.breakpoint_threshold_amount,
              buffer_size=self.params.semantic.buffer_size
          )


    def separate_files(self):
      """
      Separate Files based on extensions into TEXT and VIDEO
      """
      createDir([self.config.text_dir])
      createDir([self.config.video_dir])
      text_exts = {"pdf", "docx", "doc", "txt"}
      logger.info("File Separation Started")
      path = Path(self.config.source_dir)
      for src in path.iterdir():
          if not src.is_file():
              continue
          dest_dir = self.config.text_dir if src.suffix.lower().lstrip('.') in text_exts else self.config.video_dir
          dest = os.path.join(dest_dir,src.name)
          if not os.path.exists(dest):
              shutil.copy2(src, dest)
      logger.info("File Separation Ended")

    def merge_titles_with_body(self,elements: List[Document]) -> List[Document]:
      """
      Given a list of Document elements (from UnstructuredPDFLoader in mode="elements"),
      merge each Title element with its following body elements until the next Title.

      Assumptions:
      - elements are in document order.
      - metadata['category'] indicates element type, e.g., 'Title', 'NarrativeText', 'UncategorizedText', etc.
      - You may also consider page number: if metadata includes 'page_number', only merge within same page.
      """
      merged: List[Document] = []
      buffer_text = ""
      buffer_meta = None

      file_path = Path(elements[0].metadata.get("source", ""))
      with open(file_path,'rb') as f:
            content_bytes = f.read()
      content_hash = hashlib.md5(content_bytes).hexdigest()

      def flush_buffer():
          nonlocal buffer_text, buffer_meta
          if buffer_text.strip() and buffer_meta is not None:
              # Create a new Document with combined text and metadata from the heading or first element.
              merged.append(Document(page_content=buffer_text.strip(), metadata=buffer_meta))
          buffer_text = ""
          buffer_meta = None
      title_keywords = ["title", "heading", "section_heading", "header"]
      for elem in elements:
          meta = elem.metadata or {}
          meta['content_hash'] = content_hash
          meta['filename'] = file_path.name
          category = meta.get("category", "").lower()  # e.g., 'title', 'narrativetext', etc.
          # Normalize category check:

          is_title = any(key in category for key in title_keywords)

          if is_title:
              # If there's something buffered, flush it first
              flush_buffer()
              # Start a new buffer with the title text
              buffer_text = elem.page_content or ""
              # Keep metadata: you might want to preserve more fields; here we keep the title’s metadata,
              # but you could attach combined metadata later (e.g., include page_number etc.)
              new_meta = meta.copy()
              # Track merged element IDs
              eid = meta.get("element_id")
              if eid:
                  new_meta["merged_from"] = [eid]
              buffer_meta = new_meta

          else:
              # A non-title element: either NarrativeText or UncategorizedText, etc.
              if buffer_meta is not None:
                  # Append body text under the last heading
                  buffer_text += "\n" + (elem.page_content or "")
                  # Optionally track which element_ids are merged

                  eid = elem.metadata.get("element_id") if elem.metadata else None
                  if eid:
                      buffer_meta.setdefault("merged_from", []).append(eid)
              else:
                  # No prior title: treat this as standalone block
                  merged.append(elem)
      # Flush final buffer
      flush_buffer()
      return merged

    def _merge_small_chunks(self, chunks: List[Document], min_size: int) -> List[Document]:
      """
      Merge adjacent small chunks until reaching at least min_size characters.
      """
      merged: List[Document] = []
      buffer_text = ""
      buffer_meta = None

      for doc in chunks:
          text = doc.page_content or ""
          if not text:
              continue
          if buffer_meta is None:
              # start new buffer
              buffer_text = text
              buffer_meta = doc.metadata.copy() if doc.metadata else {}

              # track merged IDs if exist
              eid = doc.metadata.get("element_id") if doc.metadata else None
              if eid:
                  buffer_meta["merged_from"] = [eid]
          else:
              # if current buffer is smaller than min_size, append
              if len(buffer_text) < min_size:
                  buffer_text += "\n" + text
                  # merge metadata element_id

                  eid = doc.metadata.get("element_id") if doc.metadata else None
                  if eid:
                      buffer_meta.setdefault("merged_from", []).append(eid)
              else:
                  # flush existing buffer
                  merged.append(Document(page_content=buffer_text.strip(), metadata=buffer_meta))
                  # start new buffer
                  buffer_text = text
                  buffer_meta = doc.metadata.copy() if doc.metadata else {}
                  eid = doc.metadata.get("element_id") if doc.metadata else None
                  if eid:
                      buffer_meta["merged_from"] = [eid]
      # Flush last buffer
      if buffer_meta is not None and buffer_text:
          merged.append(Document(page_content=buffer_text.strip(), metadata=buffer_meta))
      return merged


    def _load_file(self, file_path: Path) -> list[Document]:
      """
      Given file path
      Load corresponding document loader,
      Return List[Document]
      """
      ext = file_path.suffix.lower()
      LOADER_MAP = {
          '.pdf': lambda p: UnstructuredPDFLoader(str(p), mode=self.params.loader.mode,strategy=self.params.loader.strategy),
          '.docx': lambda p: UnstructuredWordDocumentLoader(str(p),mode=self.params.loader.mode,strategy=self.params.loader.strategy),
          '.doc': lambda p: UnstructuredWordDocumentLoader(str(p),mode=self.params.loader.mode,strategy=self.params.loader.strategy),
          '.txt': lambda p: TextLoader(str(p), autodetect_encoding=True),
      }

      loader_fn = LOADER_MAP.get(ext)

      if not loader_fn:
          logger.warning(f"Unsupported file type: {file_path}")
          return []

      try:
          loader = loader_fn(file_path)
          docs = loader.load()
          if isinstance(docs, Document):
                docs = [docs]

          if self.params.loader.mode == "elements":
            try:
                docs = self.merge_titles_with_body(docs)
            except Exception as e:
                logger.warning(f"Title-body merging failed for {file_path}: {e}")
                # fallback: keep original docs
            # Ensure all are Document instances
          final_docs: List[Document] = []
          for d in docs:
              if isinstance(d, Document):
                  final_docs.append(d)
              else:
                  final_docs.append(Document(page_content=str(d), metadata={"source": str(file_path),'filename':Path(file_path).name}))
          return final_docs
      except Exception as e:
          logger.exception(f"Failed to load {file_path}")
          return []

    def chunk_docs(self,docs: List[Document],fname) -> List[Document]:
      final_chunks:List[Document] = []
      logger.info("Chunking Started for file {fname}")
      global_idx = 0
      for sec_doc in docs:
        sec_text = sec_doc.page_content or ""
        if not sec_text.strip():
          continue
        sub_docs = [sec_doc]
        for sub_doc in sub_docs:
          text = sub_doc.page_content or ""
          if len(text) <= self.params.recursive.chunk_size:
            chunk_meta = sub_doc.metadata.copy() if sub_doc.metadata else {}
            idx = 0
            content_hash = chunk_meta.get('content_hash',"")
            heading = chunk_meta.get("category","").replace(" ","_")
            chunk_id = f"{content_hash}_{heading}_{global_idx}"
            chunk_meta['chunk_idx'] = global_idx
            chunk_meta['chunk_id'] = chunk_id
            final_chunks.append(Document(page_content=text,metadata=chunk_meta))
            global_idx += 1
          else:
            split_docs = self.recursive_splitter.split_text(text)
            for idx,piece in enumerate(split_docs):
              piece = piece.strip()
              if not piece:
                continue
              chunk_meta = sub_doc.metadata.copy() if sub_doc.metadata else {}
              content_hash = chunk_meta.get('content_hash',"")
              heading = chunk_meta.get("category","").replace(" ","_")
              chunk_id = f"{content_hash}_{heading}_{global_idx}"
              chunk_meta['chunk_idx'] = global_idx
              chunk_meta['chunk_id'] = chunk_id
              final_chunks.append(Document(page_content=piece,metadata=chunk_meta))
              global_idx += 1
      logger.info(f"Chunking completed: {len(final_chunks)} chunks for file {fname}")
      return final_chunks


    def load_documents(self):
      self.separate_files()
      logger.info("Document loading started")

      docs = []
      chunks = []
      files = list(Path(self.config.text_dir).iterdir())
      # Parallelize document loading
      with concurrent.futures.ThreadPoolExecutor(max_workers=self.max_workers) as executor:
          for doc_list in executor.map(self._load_file, files):
              docs.append(doc_list)
      flat_docs = []
      for doc in docs:
        flat_docs.extend(doc)
        fname = ""
        if len(doc):
          fname = Path(doc[0].metadata.get('source','')).name
        chunks.extend(self.chunk_docs(doc,fname))
      logger.info(f"Document loading completed: {len(docs)} documents")
      return flat_docs,chunks
